import os
import json
from typing import Dict, List
import google.generativeai as genai
from pathlib import Path

class DocumentService:
    def __init__(self):
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        
        if self.gemini_api_key:
            genai.configure(api_key=self.gemini_api_key)
            self.gemini_model = genai.GenerativeModel('gemini-2.0-flash-exp')
    
    async def extract_text_from_pdf(self, file_path: str) -> Dict:
        """Extrait le texte d'un PDF"""
        try:
            import PyPDF2
            
            text = ""
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
            
            return {
                "success": True,
                "text": text.strip(),
                "page_count": page_count
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur extraction PDF : {str(e)}",
                "note": "Installer PyPDF2: pip install PyPDF2"
            }
    
    async def extract_text_from_docx(self, file_path: str) -> Dict:
        """Extrait le texte d'un Word"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "success": True,
                "text": text.strip(),
                "page_count": len(doc.paragraphs)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur extraction Word : {str(e)}",
                "note": "Installer python-docx: pip install python-docx"
            }
    
    async def extract_text_from_excel(self, file_path: str) -> Dict:
        """Extrait les données d'un Excel"""
        try:
            import pandas as pd
            
            # Lire toutes les feuilles
            excel_file = pd.ExcelFile(file_path)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"\n\n=== Feuille: {sheet_name} ===\n\n"
                text += df.to_string()
            
            return {
                "success": True,
                "text": text.strip(),
                "page_count": len(excel_file.sheet_names)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur extraction Excel : {str(e)}",
                "note": "Installer pandas et openpyxl: pip install pandas openpyxl"
            }
    
    async def extract_text_from_image(self, file_path: str) -> Dict:
        """Extrait le texte d'une image (OCR)"""
        try:
            from PIL import Image
            import pytesseract
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='fra+eng')
            
            return {
                "success": True,
                "text": text.strip(),
                "page_count": 1
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur OCR : {str(e)}",
                "note": "Installer Pillow et pytesseract: pip install Pillow pytesseract (+ Tesseract-OCR)"
            }
    
    async def extract_text(self, file_path: str, file_type: str) -> Dict:
        """Extrait le texte selon le type de fichier"""
        if file_type == "pdf":
            return await self.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            return await self.extract_text_from_docx(file_path)
        elif file_type in ["xlsx", "xls"]:
            return await self.extract_text_from_excel(file_path)
        elif file_type in ["jpg", "jpeg", "png", "gif", "bmp"]:
            return await self.extract_text_from_image(file_path)
        else:
            return {
                "success": False,
                "error": f"Type de fichier non supporté : {file_type}"
            }
    
    async def analyze_document_with_ai(self, text: str, filename: str) -> Dict:
        """Analyse le document avec Gemini"""
        try:
            prompt = f"""Tu es un expert en analyse de documents. Analyse ce document et fournis une analyse complète.

DOCUMENT: {filename}

CONTENU:
{text[:10000]}  # Limiter à 10000 caractères

FOURNIS une analyse JSON avec :
1. summary: Résumé du document (200 mots max)
2. key_points: Liste des 5-10 points clés
3. entities: Entités importantes (noms, dates, montants, lieux, organisations)
4. categories: Catégories du document (facture, contrat, rapport, CV, etc.)
5. sentiment: Sentiment général (positive, negative, neutral)
6. language: Langue détectée (fr, en, etc.)
7. document_type: Type précis (CV, facture, contrat, rapport, etc.)

Réponds UNIQUEMENT en JSON valide :
{{
    "summary": "...",
    "key_points": ["point 1", "point 2", ...],
    "entities": {{
        "persons": ["nom1", "nom2"],
        "dates": ["date1", "date2"],
        "amounts": ["montant1", "montant2"],
        "organizations": ["org1", "org2"],
        "locations": ["lieu1", "lieu2"]
    }},
    "categories": ["catégorie1", "catégorie2"],
    "sentiment": "positive|negative|neutral",
    "language": "fr|en|...",
    "document_type": "type"
}}"""

            response = self.gemini_model.generate_content(prompt)
            
            response_text = response.text.strip()
            if response_text.startswith("```json"):
                response_text = response_text[7:]
            if response_text.endswith("```"):
                response_text = response_text[:-3]
            
            analysis = json.loads(response_text.strip())
            
            return {
                "success": True,
                "analysis": analysis
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur analyse IA : {str(e)}"
            }
    
    async def answer_question(self, text: str, question: str) -> Dict:
        """Répond à une question sur le document"""
        try:
            prompt = f"""Tu es un assistant qui répond aux questions sur un document.

DOCUMENT:
{text[:8000]}

QUESTION: {question}

Réponds de manière précise et concise en te basant uniquement sur le contenu du document.
Si l'information n'est pas dans le document, dis-le clairement."""

            response = self.gemini_model.generate_content(prompt)
            
            return {
                "success": True,
                "answer": response.text.strip()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur réponse : {str(e)}"
            }
    
    def get_supported_formats(self) -> List[Dict]:
        """Retourne les formats supportés"""
        return [
            {
                "type": "pdf",
                "name": "PDF",
                "icon": "📄",
                "extensions": [".pdf"],
                "description": "Documents PDF"
            },
            {
                "type": "docx",
                "name": "Word",
                "icon": "📝",
                "extensions": [".docx", ".doc"],
                "description": "Documents Word"
            },
            {
                "type": "xlsx",
                "name": "Excel",
                "icon": "📊",
                "extensions": [".xlsx", ".xls"],
                "description": "Feuilles de calcul Excel"
            },
            {
                "type": "image",
                "name": "Image",
                "icon": "🖼️",
                "extensions": [".jpg", ".jpeg", ".png", ".gif", ".bmp"],
                "description": "Images (OCR)"
            }
        ]
